#!/usr/bin/env python3
"""Convert the X console manual markdown into a polished .docx with embedded screenshots."""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = os.path.expanduser("~/biz-ops/docs/x-console-manual/README.md")
OUT = os.path.expanduser("~/Downloads/X浏览控制台-运营操作手册.docx")
BASE = os.path.dirname(SRC)

doc = Document()

# ---- base styles ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for name, size, color in [("Heading 1", 16, "1F4E79"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 11.5, "404040")]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), "微软雅黑")

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

def add_runs(par, text):
    """Add runs to paragraph honoring **bold** markers."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        r = par.add_run(m.group(1)); r.bold = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])

def set_cell_shading(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def add_table(rows):
    header, body = rows[0], rows[2:]
    t = doc.add_table(rows=len(body) + 1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ""
        add_runs(c.paragraphs[0], htxt)
        for r in c.paragraphs[0].runs: r.bold = True
        set_cell_shading(c, "DCE6F1")
    for i, row in enumerate(body):
        for j, cell_text in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = ""
            add_runs(c.paragraphs[0], cell_text)
    doc.add_paragraph()

lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
first_h1_skipped = False
while i < len(lines):
    line = lines[i].rstrip()
    i += 1
    if not line.strip():
        continue
    if line.startswith("> ") and not first_h1_skipped:
        # top intro quote -> styled intro paragraph
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        add_runs(p, line[2:].strip())
        for r in p.runs: r.italic = True; r.font.color.rgb = RGBColor.from_string("595959")
        continue
    if line.startswith("# "):
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(line[2:].strip())
        r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor.from_string("1F4E79")
        doc.add_paragraph()
        continue
    if line.startswith("## 目录"):
        continue  # skip TOC (Word has its own nav)
    if line.startswith("## "):
        txt = line[3:].strip()
        m = re.match(r"^\[(.+)\]\(#.+\)$", txt)
        doc.add_heading(m.group(1) if m else txt, level=1)
        first_h1_skipped = True
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
        continue
    if line.startswith("---"):
        continue
    if line.startswith("!["):
        m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if m:
            alt, rel = m.group(1), m.group(2)
            path = os.path.join(BASE, rel)
            if os.path.exists(path):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Cm(15.5))
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rc = cap.add_run(alt); rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = RGBColor.from_string("808080")
        continue
    if line.startswith("|"):
        # collect table block
        tbl = []
        while i - 1 < len(lines) and lines[i - 1].strip().startswith("|"):
            rowline = lines[i - 1].strip()
            cells = [c.strip() for c in rowline.strip("|").split("|")]
            tbl.append(cells)
            i += 1
        if len(tbl) >= 2:
            add_table(tbl)
        continue
    m = re.match(r"^(\d+)\.\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_runs(p, m.group(2))
        continue
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, line[2:].strip())
        continue
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        add_runs(p, line[2:].strip())
        for r in p.runs: r.italic = True; r.font.color.rgb = RGBColor.from_string("595959")
        continue
    if line.startswith("*") and line.endswith("*"):
        p = doc.add_paragraph()
        r = p.add_run(line.strip("*")); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("808080")
        continue
    p = doc.add_paragraph()
    add_runs(p, line)

doc.save(OUT)
print("saved:", OUT, os.path.getsize(OUT), "bytes")
